from pdf2image import convert_from_path
#from docgen import docx_to_pdf
from docgen import localize
import docx
import io
import tempfile
from docgen import docx2pdf
from pathlib import Path
from docgen import img_tools

def create_a_table(heading, data, column_headers=None, output=None):
    """
    Args:
        heading:
        data:
        column_headers:
        output:

    Returns:
        True if output specified, else bytes
    """
    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading(heading, 0)

    # Creating a table object
    table = doc.add_table(rows=1, cols=2)

    if column_headers:
        for i, column_header in enumerate(column_headers):
            row = table.rows[0].cells
            row[i].text = column_header

    # Adding data from the list to the table
    for id, name in data:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = name

    if output:
        doc.save(output)
        return True
    else:
        with io.BytesIO() as file_stream:
            # Save the .docx to the buffer
            doc.save(file_stream)
            # Reset the buffer's file-pointer to the beginning of the file
            bytes_model = file_stream.getvalue()
            return bytes_model

"""
    # Replace text/fonts
    # Replace text with HWR
    # Install PIL for Windows
    # Test on RHEL
    # Check coordinates against rendered image
"""


def main(root):

    # Table data in a form of list
    data = (
        ("Taylor", 'Some Text'),
        (2, 'Additional Text'),
        (3, 'Also this')
    )
    data = []
    heading = 'TAYLOR'
    column_headers=["TAYLOR", "TAYLOR"]

    if False:
        docx_temp = tempfile.TemporaryFile()
        pdf_temp = tempfile.TemporaryFile()
        img_path = tempfile.TemporaryFile()
    else:
        docx_temp = Path(f"./temp/{root}.docx")
        pdf_temp = Path(f"./temp/{root}.pdf")
        img_path = f"./temp/{root}" + "{}.jpg"

        pdf_temp = r"C:\Users\tarchibald\Downloads\EXAMPLES\french_census_0000.pdf"

    if False:
        success = create_a_table(heading, data, column_headers, docx_temp)

    if False:
        # Convert to PDF and Image
        docx2pdf.convert(docx_temp, pdf_temp)

    # Localize
    localization = localize.generate_localization_from_file(pdf_temp)
    text = "".join([t["text"] for t in localization[0]["textboxes_character"]])

    print("done with localization")

    # Save as image
    print(img_path)
    img_tools.convert_pdf_to_img_paths(pdf_temp, img_path)

if __name__ == '__main__':
    main("TEMPLATE")