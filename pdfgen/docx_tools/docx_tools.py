import random

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ROW_HEIGHT_RULE
import docx
from typing import Literal
from docx.shared import Inches
import numpy as np
from docx.oxml import table
from docx import Document
# Importing the necessary functions to set the row height
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def new_document():
    """ Default: 200 DPI

    Returns:

    """
    document = Document()
    section = document.sections[0]
    section.page_height = Inches(15)
    section.page_width = Inches(20)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)
    section.top_margin = Inches(.25)
    section.bottom_margin = Inches(.25)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    return document

def random_document():
    document = Document()
    section = document.sections[0]
    width = random.randint(10,20)
    section.page_width = Inches(width)
    section.page_height = Inches(width+random.randint(0,5))
    section.left_margin = Inches(.1)
    section.right_margin = Inches(.1)
    section.top_margin = Inches(.1)
    section.bottom_margin = Inches(.1)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    return document

def landscape(section):
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

def create_document_base(heading=None,
                         table_data=None,
                         output=None,
                         page_orientation: Literal['portrait', 'landscape'] = "landscape",
                         column_autofit=False,
                         grid=True,
                         widths=None,
                         expand_row_height=True,
                         no_breaking_spaces_hack=True):
    doc = random_document()


    if page_orientation=="landscape":
        if doc.sections[-1].page_height > doc.sections[-1].page_width:
            landscape(doc.sections[-1])

    # Add a Title to the document
    if heading:
        doc.add_heading(heading, 0)


    if table_data:
        table = doc.add_table(rows=0, cols=len(table_data[0]))
        if not column_autofit:
            table.autofit = False
            table.allow_autofit = False
        else:
            table.autofit = True
            table.allow_autofit = True # toothless?

        if not widths is None:
            widths = np.asarray(widths)

        for i, row_data in enumerate(table_data):
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            for j, cell in enumerate(row):
                txt = str(row_data[j])
                if no_breaking_spaces_hack:
                    txt = txt.replace(" ","\u00A0").replace("-","\u2011")
                row[j].text = str(txt)
                if not widths is None:
                    if widths.ndim == 1: # columns have uniform width
                        table.cell(i, j).width = Inches(widths[j])
                    else: # cells have unique widths
                        table.cell(i, j).width = Inches(widths[i,j])
                        #table.gridspan
                        table.cell(i,j).merge(table.cell(0,j))

            # Don't expand height
            if not expand_row_height:
                table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                table.rows[i].height = Inches(.25)

    if grid:
        table.style = 'Table Grid'

    if output:
        doc.save(output)

    return doc

def addCheckbox(para, box_id, name, checked):

    run = para.add_run()
    tag = run._r
    fldchar = docx.oxml.shared.OxmlElement('w:fldChar')
    fldchar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')

    ffdata = docx.oxml.shared.OxmlElement('w:ffData')
    name = docx.oxml.shared.OxmlElement('w:name')
    name.set(docx.oxml.ns.qn('w:val'), cb_name)
    enabled = docx.oxml.shared.OxmlElement('w:enabled')
    calconexit = docx.oxml.shared.OxmlElement('w:calcOnExit')
    calconexit.set(docx.oxml.ns.qn('w:val'), '0')

    checkbox = docx.oxml.shared.OxmlElement('w:checkBox')
    sizeauto = docx.oxml.shared.OxmlElement('w:sizeAuto')
    default = docx.oxml.shared.OxmlElement('w:default')

    if checked:
        default.set(docx.oxml.ns.qn('w:val'), '1')
    else:
        default.set(docx.oxml.ns.qn('w:val'), '0')

    checkbox.append(sizeauto)
    checkbox.append(default)
    ffdata.append(name)
    ffdata.append(enabled)
    ffdata.append(calconexit)
    ffdata.append(checkbox)
    fldchar.append(ffdata)
    tag.append(fldchar)

    run2 = para.add_run()
    tag2 = run2._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), str(box_id))
    start.set(docx.oxml.ns.qn('w:name'), name)
    tag2.append(start)

    run3 = para.add_run()
    tag3 = run3._r
    instr = docx.oxml.OxmlElement('w:instrText')
    instr.text = 'FORMCHECKBOX'
    tag3.append(instr)

    run4 = para.add_run()
    tag4 = run4._r
    fld2 = docx.oxml.shared.OxmlElement('w:fldChar')
    fld2.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
    tag4.append(fld2)

    run5 = para.add_run()
    tag5 = run5._r
    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), str(box_id))
    end.set(docx.oxml.ns.qn('w:name'), name)
    tag5.append(end)

    return


def add_picture():
    for i in range(5):
        pic_path = f"Table_Images\pic_{i}.jpg"

        cell = cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell_p, cell_f, cell_r = paragraph_format_run(cell)

        cell_r.add_picture(pic_path, width=Inches(1.25))

def setRowHeight(row, height):
    """
    Sets the height of a table row.

    `Row` is a `docx.table._Row` object. `height` is the desired height in EMU.
    """
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height.twips))

    trPr = OxmlElement('w:trPr')
    trPr.append(trHeight)

    row._tr.append(trPr)

def uneven_columns(row, value=2):
    """
    Sets the height of a table row.

    `Row` is a `docx.table._Row` object. `height` is the desired height in EMU.
    <w:gridSpan w:val="2"/>
    """

    gridspan = OxmlElement('w:gridSpan')
    gridspan.set(qn('w:val'), str(value))

    tcPr = OxmlElement('w:tcPr')
    tcPr.append(gridspan)
    row._tc.append(tcPr)


"""
 Can you add pictures and a note about what text they have?
 Precompute and resize lines
 
"""