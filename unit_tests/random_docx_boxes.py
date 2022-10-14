from pdf2image import convert_from_path
#from docgen import docx_to_pdf
from docgen import localize
import docx
import io
import tempfile
from pathlib import Path
from docx.shared import Cm, Inches
import random
from docgen.utils import file_incrementer

def create_docx(output="test.docx", heading = 'TAYLOR'):
    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading(heading, 0)

    # Creating a table object
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Table Grid'

    for i in range(10):
        w = 3*random.random()
        table.rows[i].cells[0].width = Inches(w)
        print(w)

    doc.save(file_incrementer(output))

data = {

}

if __name__ == '__main__':
    create_docx()

